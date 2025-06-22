#!/usr/bin/env python3
"""
PDF to Word Converter using Gemini API
This script converts PDF files to Word documents using Google's Gemini API
for intelligent text processing and formatting.
"""

import os
import sys
import argparse
from pathlib import Path
import google.generativeai as genai
from docx import Document
from docx.shared import Inches
import PyPDF2
import pdfplumber
import fitz  # PyMuPDF
import io
import re
import time
from typing import Optional, List, Tuple
import logging

# Import advanced features
try:
    from advanced_features import ConversionStats, DocumentAnalyzer, QualityChecker, create_conversion_report
except ImportError:
    # Fallback if advanced features not available
    ConversionStats = None
    DocumentAnalyzer = None
    QualityChecker = None
    create_conversion_report = None

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class PDFToWordConverter:
    def __init__(self, api_key: str, preferred_model: str = None, enable_stats: bool = True):
        """
        Initialize the converter with Gemini API key

        Args:
            api_key (str): Google Gemini API key
            preferred_model (str, optional): Preferred model to use
            enable_stats (bool): Enable conversion statistics tracking
        """
        self.api_key = api_key
        genai.configure(api_key=api_key)

        # Initialize statistics tracking
        self.stats = ConversionStats() if ConversionStats and enable_stats else None
        self.extraction_method_used = None

        # Get all available models dynamically
        self.available_models = self._fetch_available_models()

        # Set model preference order
        if preferred_model and preferred_model in self.available_models:
            self.model_names = [preferred_model] + [m for m in self.available_models if m != preferred_model]
        else:
            # Default fallback order with latest models first
            self.model_names = self.available_models

        self.current_model = None
        self.current_model_name = None
        self._initialize_model()

    def _fetch_available_models(self):
        """Fetch all available Gemini models dynamically"""
        try:
            # Get all available models
            models = []
            for model in genai.list_models():
                if hasattr(model, 'name') and 'gemini' in model.name.lower():
                    model_name = model.name.replace('models/', '')
                    models.append(model_name)

            if not models:
                # Fallback to known models if API call fails
                models = [
                    'gemini-2.0-flash-exp',
                    'gemini-1.5-flash',
                    'gemini-1.5-pro',
                    'gemini-pro'
                ]

            # Sort models with latest first
            models.sort(key=lambda x: (
                '2.0' in x,  # Gemini 2.0 models first
                'flash' in x,  # Flash models preferred
                'exp' in x,  # Experimental versions
                x
            ), reverse=True)

            logger.info(f"Found {len(models)} available Gemini models")
            return models

        except Exception as e:
            logger.warning(f"Failed to fetch available models: {str(e)}")
            # Return default fallback models
            return [
                'gemini-2.0-flash-exp',
                'gemini-1.5-flash',
                'gemini-1.5-pro',
                'gemini-pro'
            ]

    def _initialize_model(self):
        """Initialize the best available model with fallback"""
        for model_name in self.model_names:
            try:
                self.current_model = genai.GenerativeModel(model_name)
                self.current_model_name = model_name
                logger.info(f"Successfully initialized model: {model_name}")
                break
            except Exception as e:
                logger.warning(f"Failed to initialize {model_name}: {str(e)}")
                continue

        if self.current_model is None:
            raise Exception("Failed to initialize any Gemini model")

    def get_available_models(self):
        """Get list of all available models"""
        return self.available_models

    def get_current_model(self):
        """Get currently active model name"""
        return self.current_model_name

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text content from PDF file using multiple methods for best results

        Args:
            pdf_path (str): Path to the PDF file

        Returns:
            str: Extracted text content with proper spacing
        """
        extraction_methods = [
            ("pdfplumber", self._extract_with_pdfplumber),
            ("PyMuPDF", self._extract_with_pymupdf),
            ("PyPDF2", self._extract_with_pypdf2)
        ]

        for method_name, extraction_func in extraction_methods:
            try:
                logger.info(f"Trying extraction with {method_name}...")
                text_content = extraction_func(pdf_path)

                if text_content and text_content.strip():
                    logger.info(f"Successfully extracted text using {method_name}")
                    self.extraction_method_used = method_name

                    # Apply preprocessing to fix spacing issues
                    processed_text = self._preprocess_spacing(text_content)
                    return processed_text
                else:
                    logger.warning(f"{method_name} returned empty text")

            except Exception as e:
                logger.warning(f"{method_name} extraction failed: {str(e)}")
                continue

        raise Exception("All PDF extraction methods failed")

    def _extract_with_pdfplumber(self, pdf_path: str) -> str:
        """Extract text using pdfplumber (best for spacing)"""
        text_content = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n\n"
            logger.info(f"pdfplumber: Extracted text from {len(pdf.pages)} pages")
        return text_content

    def _extract_with_pymupdf(self, pdf_path: str) -> str:
        """Extract text using PyMuPDF (good for complex layouts)"""
        text_content = ""
        doc = fitz.open(pdf_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_text = page.get_text()
            if page_text:
                text_content += page_text + "\n\n"
        doc.close()
        logger.info(f"PyMuPDF: Extracted text from {len(doc)} pages")
        return text_content

    def _extract_with_pypdf2(self, pdf_path: str) -> str:
        """Extract text using PyPDF2 (fallback method)"""
        text_content = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n\n"
            logger.info(f"PyPDF2: Extracted text from {len(pdf_reader.pages)} pages")
        return text_content

    def _preprocess_spacing(self, text: str) -> str:
        """
        Intelligent preprocessing to fix common spacing issues before AI processing
        """
        logger.info("Applying intelligent spacing preprocessing...")

        # Fix merged words with common patterns
        processed_text = text

        # Pattern 1: Add space before capital letters in merged words (camelCase)
        # But avoid breaking acronyms or proper nouns
        processed_text = re.sub(r'([a-z])([A-Z])', r'\1 \2', processed_text)

        # Pattern 2: Add space between word and number
        processed_text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', processed_text)
        processed_text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', processed_text)

        # Pattern 3: Fix common merged words (add more as needed)
        common_merges = {
            r'andthe': 'and the',
            r'ofthe': 'of the',
            r'inthe': 'in the',
            r'tothe': 'to the',
            r'forthe': 'for the',
            r'withthe': 'with the',
            r'onthe': 'on the',
            r'atthe': 'at the',
            r'bythe': 'by the',
            r'fromthe': 'from the',
            r'thatthe': 'that the',
            r'thisthe': 'this the',
            r'itis': 'it is',
            r'thisis': 'this is',
            r'thereis': 'there is',
            r'therefor': 'therefore',
            r'however': 'however',
            r'moreover': 'moreover',
            r'furthermore': 'furthermore',
        }

        for merged, separated in common_merges.items():
            processed_text = re.sub(merged, separated, processed_text, flags=re.IGNORECASE)

        # Pattern 4: Fix punctuation spacing
        processed_text = re.sub(r'([.!?])([A-Z])', r'\1 \2', processed_text)
        processed_text = re.sub(r'([,;:])([a-zA-Z])', r'\1 \2', processed_text)

        # Pattern 5: Clean up multiple spaces
        processed_text = re.sub(r' +', ' ', processed_text)

        # Pattern 6: Fix line breaks that split words
        processed_text = re.sub(r'([a-z])-\n([a-z])', r'\1\2', processed_text)

        logger.info("Spacing preprocessing completed")
        return processed_text

    def process_with_gemini(self, text_content: str) -> str:
        """
        Process extracted text with Gemini API for better formatting with model fallback

        Args:
            text_content (str): Raw text extracted from PDF

        Returns:
            str: Processed and formatted text
        """
        prompt = f"""
You are an expert text processor. Clean up this PDF-extracted text with EXTREME attention to word spacing:

CRITICAL SPACING RULES:
1. **MERGED WORDS**: Look for words stuck together and separate them properly
   - Example: "thequick" → "the quick"
   - Example: "andthe" → "and the"
   - Example: "itis" → "it is"
   - Example: "therefor" → "therefore"

2. **MISSING SPACES**: Add spaces where words are clearly merged
   - Between articles and nouns: "thedog" → "the dog"
   - Between prepositions: "inthe" → "in the"
   - Before capital letters: "wordAnother" → "word Another"
   - Between numbers and words: "5years" → "5 years"

3. **PRESERVE CORRECT SPACING**: Don't add extra spaces where they already exist correctly

4. **OTHER FIXES**:
   - Fix obvious OCR errors
   - Maintain paragraph structure
   - Fix punctuation spacing: "word.Another" → "word. Another"
   - Remove random line breaks within sentences
   - Keep original meaning and structure

5. **DO NOT**:
   - Add titles or headers
   - Change the document structure
   - Add commentary or explanations

EXAMPLES OF SPACING FIXES:
- "Thisisanimportantdocument" → "This is an important document"
- "Thecompanyreported5milliondollars" → "The company reported 5 million dollars"
- "However,theresults" → "However, the results"

Text to process:
{text_content}

Return ONLY the corrected text with proper spacing:"""

        # Try models in fallback order
        for i, model_name in enumerate(self.model_names):
            try:
                model = genai.GenerativeModel(model_name)
                response = model.generate_content(prompt)
                logger.info(f"Text processed successfully with {model_name}")
                return response.text

            except Exception as e:
                logger.warning(f"Failed to process with {model_name}: {str(e)}")
                if i < len(self.model_names) - 1:
                    logger.info(f"Falling back to next model...")
                    continue
                else:
                    logger.error("All models failed, returning original text")
                    return text_content

        return text_content

    def _post_process_spacing(self, text: str) -> str:
        """
        Final validation and fixing of spacing issues after AI processing
        """
        logger.info("Applying post-processing spacing validation...")

        processed_text = text

        # Check for remaining merged words using common patterns
        suspicious_patterns = [
            r'[a-z][A-Z]',  # camelCase
            r'[a-zA-Z]\d',  # word followed by number
            r'\d[a-zA-Z]',  # number followed by word
            r'[.!?][A-Z]',  # punctuation followed by capital
            r'[,;:][a-zA-Z]',  # punctuation followed by letter
        ]

        issues_found = 0
        for pattern in suspicious_patterns:
            matches = re.findall(pattern, processed_text)
            if matches:
                issues_found += len(matches)

        if issues_found > 0:
            logger.warning(f"Found {issues_found} potential spacing issues in post-processing")

            # Apply final fixes
            processed_text = re.sub(r'([a-z])([A-Z])', r'\1 \2', processed_text)
            processed_text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', processed_text)
            processed_text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', processed_text)
            processed_text = re.sub(r'([.!?])([A-Z])', r'\1 \2', processed_text)
            processed_text = re.sub(r'([,;:])([a-zA-Z])', r'\1 \2', processed_text)

            # Clean up multiple spaces
            processed_text = re.sub(r' +', ' ', processed_text)

            logger.info("Applied final spacing corrections")
        else:
            logger.info("No spacing issues detected in post-processing")

        return processed_text

    def _detect_heading_level(self, text: str) -> int:
        """
        Intelligently detect heading level based on text patterns
        """
        text = text.strip()

        # Level 1: Major headings
        if (text.isupper() and len(text) < 60) or \
           any(text.startswith(prefix) for prefix in ['CHAPTER', 'PART', 'SECTION I', 'APPENDIX']):
            return 1

        # Level 2: Secondary headings
        if text.startswith(('Chapter', 'Section', 'Part')) or \
           (text.endswith(':') and len(text) < 50) or \
           any(text.startswith(f'{i}.') for i in range(1, 21)):
            return 2

        # Level 3: Sub-headings
        if any(text.startswith(f'{i}.{j}') for i in range(1, 11) for j in range(1, 11)) or \
           text.startswith(('Introduction', 'Conclusion', 'Summary', 'Overview')):
            return 3

        return 0  # Not a heading

    def _is_heading(self, text: str) -> bool:
        """
        Enhanced heading detection with multiple criteria
        """
        text = text.strip()

        # Skip very long lines
        if len(text) > 100:
            return False

        # Check various heading patterns
        heading_patterns = [
            text.isupper() and len(text) < 60,  # All caps, reasonable length
            text.endswith(':') and len(text) < 50,  # Ends with colon
            any(text.startswith(prefix) for prefix in [
                'CHAPTER', 'PART', 'SECTION', 'APPENDIX', 'INTRODUCTION', 'CONCLUSION',
                'Chapter', 'Part', 'Section', 'Appendix', 'Introduction', 'Conclusion'
            ]),
            any(text.startswith(f'{i}.') for i in range(1, 21)),  # Numbered 1-20
            any(text.startswith(f'{i}.{j}') for i in range(1, 11) for j in range(1, 11)),  # Sub-numbered
            text.startswith(('Summary', 'Overview', 'Abstract', 'References', 'Bibliography'))
        ]

        return any(heading_patterns)

    def create_word_document(self, processed_text: str, output_path: str) -> None:
        """
        Create a Word document with enhanced formatting and structure

        Args:
            processed_text (str): Formatted text content
            output_path (str): Path for the output Word document
        """
        try:
            doc = Document()

            # Enhanced document structure analysis
            lines = processed_text.split('\n')

            # First pass: identify document structure
            headings = []
            for i, line in enumerate(lines):
                if line.strip() and self._is_heading(line):
                    level = self._detect_heading_level(line)
                    headings.append((i, line.strip(), level))

            # Second pass: build document with proper structure
            i = 0
            while i < len(lines):
                line = lines[i].strip()

                if not line:
                    # Add spacing for empty lines
                    doc.add_paragraph()
                    i += 1
                    continue

                # Check if this line is a heading
                is_heading_line = False
                heading_level = 0

                for heading_idx, heading_text, level in headings:
                    if heading_idx == i:
                        is_heading_line = True
                        heading_level = level
                        break

                if is_heading_line:
                    # Add heading with appropriate level
                    doc.add_heading(line, level=max(1, heading_level))
                else:
                    # Regular paragraph - collect consecutive lines
                    paragraph_lines = [line]
                    j = i + 1

                    # Collect lines until next heading or empty line
                    while j < len(lines):
                        next_line = lines[j].strip()
                        if not next_line:  # Empty line ends paragraph
                            break

                        # Check if next line is a heading
                        is_next_heading = any(heading_idx == j for heading_idx, _, _ in headings)
                        if is_next_heading:
                            break

                        paragraph_lines.append(next_line)
                        j += 1

                    # Create paragraph with proper spacing
                    if paragraph_lines:
                        paragraph_text = ' '.join(paragraph_lines)
                        paragraph = doc.add_paragraph()
                        paragraph.add_run(paragraph_text)

                    i = j - 1  # Adjust index for collected lines

                i += 1

            # Save the document
            doc.save(output_path)
            logger.info(f"Word document saved with enhanced formatting: {output_path}")

        except Exception as e:
            logger.error(f"Error creating Word document: {str(e)}")
            raise
    
    def convert_pdf_to_word(self, pdf_path: str, output_path: Optional[str] = None,
                           generate_report: bool = False) -> str:
        """
        Enhanced main conversion method with analytics and quality checking

        Args:
            pdf_path (str): Path to input PDF file
            output_path (str, optional): Path for output Word file
            generate_report (bool): Generate detailed conversion report

        Returns:
            str: Path to the created Word document
        """
        start_time = time.time()
        success = False
        pages_processed = 0

        try:
            # Validate input file
            if not os.path.exists(pdf_path):
                raise FileNotFoundError(f"PDF file not found: {pdf_path}")

            # Generate output path if not provided
            if output_path is None:
                pdf_name = Path(pdf_path).stem
                output_path = f"{pdf_name}_converted.docx"

            logger.info(f"Starting enhanced conversion: {pdf_path} -> {output_path}")

            # Step 1: Extract text from PDF
            logger.info("Extracting text from PDF...")
            raw_text = self.extract_text_from_pdf(pdf_path)

            if not raw_text.strip():
                raise ValueError("No text content found in PDF")

            # Estimate pages processed (rough estimate based on text length)
            pages_processed = max(1, len(raw_text) // 2000)  # ~2000 chars per page

            # Step 2: Process with Gemini API
            logger.info("Processing text with Gemini API...")
            ai_processed_text = self.process_with_gemini(raw_text)

            # Step 3: Post-process for final spacing validation
            logger.info("Applying final spacing validation...")
            final_text = self._post_process_spacing(ai_processed_text)

            # Step 4: Quality checking (if available)
            quality_issues = []
            if QualityChecker:
                logger.info("Performing quality checks...")
                spacing_issues = QualityChecker.check_spacing_quality(final_text)
                text_issues = QualityChecker.check_text_quality(final_text)
                quality_issues.extend(spacing_issues + text_issues)

                if quality_issues:
                    logger.warning(f"Quality issues detected: {len(quality_issues)}")
                    for issue in quality_issues:
                        logger.warning(f"  - {issue}")
                else:
                    logger.info("✅ No quality issues detected")

            # Step 5: Create Word document
            logger.info("Creating Word document...")
            self.create_word_document(final_text, output_path)

            # Step 6: Generate report if requested
            if generate_report and DocumentAnalyzer and create_conversion_report:
                logger.info("Generating conversion report...")
                analysis = DocumentAnalyzer.analyze_text_structure(final_text)
                report = create_conversion_report(pdf_path, output_path, self.stats, analysis, quality_issues)

                report_path = output_path.replace('.docx', '_report.md')
                with open(report_path, 'w', encoding='utf-8') as f:
                    f.write(report)
                logger.info(f"Report saved to: {report_path}")

            success = True
            processing_time = time.time() - start_time
            logger.info(f"Conversion completed successfully in {processing_time:.2f}s!")

            return output_path

        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"Conversion failed after {processing_time:.2f}s: {str(e)}")
            raise

        finally:
            # Record statistics
            if self.stats:
                processing_time = time.time() - start_time
                self.stats.record_conversion(
                    success=success,
                    pages=pages_processed,
                    processing_time=processing_time,
                    model_used=self.current_model_name,
                    extraction_method=self.extraction_method_used
                )

def main():
    """Enhanced main function with advanced features"""
    parser = argparse.ArgumentParser(description='Convert PDF to Word using Gemini API with advanced features')
    parser.add_argument('pdf_file', help='Path to the PDF file to convert')
    parser.add_argument('-o', '--output', help='Output Word file path (optional)')
    parser.add_argument('-k', '--api-key', help='Gemini API key (or set GEMINI_API_KEY env var)')
    parser.add_argument('-m', '--model', help='Preferred Gemini model to use')
    parser.add_argument('-r', '--report', action='store_true', help='Generate detailed conversion report')
    parser.add_argument('--list-models', action='store_true', help='List available Gemini models')
    parser.add_argument('--stats', action='store_true', help='Show conversion statistics')
    parser.add_argument('--no-stats', action='store_true', help='Disable statistics tracking')

    args = parser.parse_args()

    # Get API key from argument or environment variable
    api_key = args.api_key or os.getenv('GEMINI_API_KEY')

    if not api_key and not args.list_models:
        logger.error("Gemini API key is required. Provide it via --api-key argument or GEMINI_API_KEY environment variable")
        sys.exit(1)

    try:
        # Initialize converter
        enable_stats = not args.no_stats
        converter = PDFToWordConverter(api_key, preferred_model=args.model, enable_stats=enable_stats)

        # Handle list models command
        if args.list_models:
            print("🤖 Available Gemini Models:")
            for i, model in enumerate(converter.get_available_models(), 1):
                current = " (current)" if model == converter.get_current_model() else ""
                print(f"  {i}. {model}{current}")
            return

        # Handle stats command
        if args.stats and converter.stats:
            print(converter.stats.get_summary())
            return

        # Perform conversion
        print(f"🚀 Starting conversion with model: {converter.get_current_model()}")
        output_file = converter.convert_pdf_to_word(args.pdf_file, args.output, generate_report=args.report)

        print(f"✅ Conversion successful!")
        print(f"📄 Input: {args.pdf_file}")
        print(f"📝 Output: {output_file}")
        print(f"🤖 Model used: {converter.get_current_model()}")
        print(f"🔧 Extraction method: {converter.extraction_method_used}")

        if args.report:
            report_file = output_file.replace('.docx', '_report.md')
            print(f"📊 Report: {report_file}")

        # Show quick stats
        if converter.stats:
            print(f"📈 Success rate: {converter.stats.get_success_rate():.1f}%")

    except Exception as e:
        logger.error(f"Conversion failed: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
